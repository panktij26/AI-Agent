"""
conftest.py
Shared pytest fixtures used across the test suite.

Key idea: none of these tests call the real Claude API. Wherever the pipeline
would call Claude, we monkeypatch/mock it with a fake, deterministic response.
This means the whole test suite runs for free, offline, in a couple of seconds,
and doesn't require ANTHROPIC_API_KEY to be set.
"""

import json
import os

import pytest


@pytest.fixture
def sample_jd_text():
    return (
        "Job Title: Backend Python Developer\n"
        "Required Skills: Python, REST APIs, PostgreSQL, Docker, AWS.\n"
    )


@pytest.fixture
def temp_txt_resume(tmp_path):
    """A plain-text resume fixture."""
    content = "Jane Doe\nPython, FastAPI, PostgreSQL, Docker, AWS\n5 years experience.\n"
    path = tmp_path / "resume.txt"
    path.write_text(content, encoding="utf-8")
    return str(path), content


@pytest.fixture
def temp_pdf_resume(tmp_path):
    """A real, minimal PDF resume fixture built with reportlab."""
    from reportlab.pdfgen import canvas

    path = tmp_path / "resume.pdf"
    known_phrase = "PDF CANDIDATE PYTHON DOCKER POSTGRESQL"
    c = canvas.Canvas(str(path))
    c.drawString(100, 750, "John Smith")
    c.drawString(100, 730, known_phrase)
    c.save()
    return str(path), known_phrase


@pytest.fixture
def temp_docx_resume(tmp_path):
    """A real, minimal DOCX resume fixture built with python-docx."""
    import docx

    path = tmp_path / "resume.docx"
    known_phrase = "DOCX CANDIDATE FASTAPI AWS KUBERNETES"
    document = docx.Document()
    document.add_paragraph("Alex Lee")
    document.add_paragraph(known_phrase)
    document.save(str(path))
    return str(path), known_phrase


@pytest.fixture
def temp_resumes_dir(tmp_path):
    """A folder with a few small, deterministic .txt resumes for pipeline tests."""
    resumes_dir = tmp_path / "resumes"
    resumes_dir.mkdir()

    (resumes_dir / "strong_fit.txt").write_text(
        "Priya\nPython, FastAPI, PostgreSQL, Docker, AWS, REST APIs, 5 years experience.\n"
    )
    (resumes_dir / "weak_fit.txt").write_text(
        "Sam\nPhotography, graphic design, Adobe Photoshop, event planning.\n"
    )
    (resumes_dir / "empty.txt").write_text("")  # edge case: empty file
    (resumes_dir / "notes.rtf").write_text("this should be ignored, unsupported type")

    return str(resumes_dir)


@pytest.fixture
def temp_jd_file(tmp_path, sample_jd_text):
    path = tmp_path / "jd.txt"
    path.write_text(sample_jd_text)
    return str(path)


def make_fake_claude_response(payload: dict):
    """
    Builds an object shaped like anthropic's Message response, containing a
    single text block whose text is the JSON-encoded payload. Used to mock
    client.messages.create(...) without hitting the real API.
    """
    class FakeBlock:
        type = "text"
        text = json.dumps(payload)

    class FakeResponse:
        content = [FakeBlock()]

    return FakeResponse()
